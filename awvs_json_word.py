#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Acunetix/Invicti 漏洞扫描 JSON 报告生成器（一键生成 Word 报告）

用法：
    python awvs_json_word.py <report.json> [output.docx]

参数：
    report.json  - Acunetix/Invicti 导出的漏洞扫描 JSON 文件路径（必须）
    output.docx  - 输出 Word 文件路径（可选，默认自动生成）
"""

import base64
import gzip
import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib.pyplot as plt
import matplotlib
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

TOOL_NAME = "Acunetix"
DEFAULT_TOOL_VERSION = "unknown"

# ========== Acunetix/Invicti 数据解析与转换 ==========
# severity 数字 -> 等级字符串
SEVERITY_MAP = {4: 'critical', 3: 'high', 2: 'medium', 1: 'low', 0: 'info'}

# 漏洞类型 -> 中文分类
TYPE_ZH_MAP = {
    'configuration': '安全配置',
    'weakcrypto': '弱加密算法',
    'sensitivedatanotoverssl': '敏感数据传输未加密',
    'informationdisclosure': '信息泄露',
    'authentication': '认证安全',
    'xss': '跨站脚本(XSS)',
    'sql': 'SQL注入',
    'defaultcontent': '默认内容',
}

# 常见漏洞英文名称 -> 中文名称/描述/修复建议（未收录的漏洞回退为原始英文）
VULN_ZH_MAP = {
    "HTTP Strict Transport Security (HSTS) Policy Not Enabled": {
        "title_zh": "未启用 HTTP 严格传输安全（HSTS）策略",
        "description_zh": "检测到 Web 应用未在响应头中实现 HTTP 严格传输安全（HSTS），即响应中缺少 Strict-Transport-Security 头。HSTS 可告知浏览器仅通过 HTTPS 访问网站，能够有效预防和缓解部分中间人（MitM）攻击。",
        "repair_zh": "建议在 Web 应用中启用 HTTP 严格传输安全（HSTS），在响应头中添加 Strict-Transport-Security 头，例如：Strict-Transport-Security: max-age=31536000; includeSubDomains。",
    },
    "Misconfigured Access-Control-Allow-Origin Header": {
        "title_zh": "CORS 跨域配置错误（Access-Control-Allow-Origin 头配置不当）",
        "description_zh": "应用未正确校验 Origin 请求头，并返回 Access-Control-Allow-Origin 且 Access-Control-Allow-Credentials 为 true。在此配置下，任意网站都可以携带用户凭证发起跨域请求并读取响应内容，等同于禁用了同源策略，第三方网站可与应用进行双向交互。",
        "repair_zh": "仅在 Access-Control-Allow-Origin 响应头中允许受信任的指定域名，切勿直接回显任意 Origin 值；如无必要，请勿设置 Access-Control-Allow-Credentials: true。",
    },
    "TLS/SSL Sweet32 attack": {
        "title_zh": "TLS/SSL Sweet32 攻击漏洞",
        "description_zh": "Sweet32 攻击是针对使用 64 位分组密码算法的 SSL/TLS 连接的漏洞，攻击者可利用生日攻击破解 HTTPS 加密会话，从而拦截或篡改加密通信内容。",
        "repair_zh": "重新配置受影响的 SSL/TLS 服务器，禁用 3DES 等 64 位分组密码套件，推荐使用 AES-128/256-GCM 等强加密套件。",
    },
    "TLS/SSL Weak Cipher Suites": {
        "title_zh": "TLS/SSL 弱加密套件",
        "description_zh": "远程主机支持包含弱或不安全属性的 TLS/SSL 密码套件，例如 CBC 模式、3DES 等旧式算法，攻击者可能利用这些弱算法实施降级攻击或破解加密通信。",
        "repair_zh": "重新配置应用，仅启用安全的密码套件，禁用 CBC 模式、3DES 等弱算法，优先使用 TLS 1.2/1.3 及 AES-GCM/CHACHA20 套件。",
    },
    "Cookies Not Marked as Secure": {
        "title_zh": "Cookie 未设置 Secure 标志",
        "description_zh": "检测到部分 Cookie 未设置 Secure 标志。设置 Secure 标志后，浏览器将仅通过加密的 SSL/TLS 信道传输该 Cookie，这对于会话 Cookie 是重要的安全保护措施；否则 Cookie 可能通过明文信道传输而被窃取。",
        "repair_zh": "为相关 Cookie 设置 Secure 标志，确保 Cookie 仅通过 HTTPS 安全信道传输。",
    },
    "Cookies with missing, inconsistent or contradictory properties": {
        "title_zh": "Cookie 属性缺失、不一致或相互矛盾",
        "description_zh": "检测到至少一个 Cookie 的属性缺失、与其它属性不一致或相互矛盾，导致 Cookie 无法被浏览器正确存储或提交。该问题本身虽不直接构成漏洞，但可能导致应用出现异常行为，进而引发二次安全问题，例如 Cookie 缺少 SameSite 属性。",
        "repair_zh": "确保 Cookie 配置符合相关标准，建议为 Cookie 设置完整且一致的属性，包括 Secure、HttpOnly 及适当的 SameSite（Strict/Lax/None）标志。",
    },
    "Content Security Policy (CSP) Not Implemented": {
        "title_zh": "未实施内容安全策略（CSP）",
        "description_zh": "检测到 Web 应用未实现内容安全策略（CSP），响应中缺少 Content-Security-Policy 头。CSP 是帮助检测和缓解跨站脚本（XSS）及数据注入攻击的重要安全层，可防止恶意内容注入和点击劫持等攻击。",
        "repair_zh": "建议在 Web 应用中实施内容安全策略，在页面响应中添加 Content-Security-Policy 头，明确列出允许加载各类资源（脚本、样式、图片等）的来源。",
    },
    "Permissions-Policy header not implemented": {
        "title_zh": "未实施 Permissions-Policy 响应头",
        "description_zh": "检测到 Web 应用未实现 Permissions-Policy 头。Permissions-Policy 允许开发者有选择地启用或禁用浏览器各项功能与 API（如摄像头、麦克风、地理位置等），从而减少攻击面。",
        "repair_zh": "在 Web 应用响应中添加 Permissions-Policy 头，按需限制浏览器功能的使用。",
    },
    "Web Application Firewall Detected": {
        "title_zh": "检测到 Web 应用防火墙（WAF）",
        "description_zh": "检测到目标服务器受 IPS/IDS/WAF 保护。扫描器在发送各种恶意载荷时，响应码、响应头或响应体发生变化，由此判定 WAF 存在。受 WAF 保护的站点扫描结果可能不完整，且大量攻击载荷可能导致扫描源 IP 被封锁。",
        "repair_zh": "如条件允许，建议对未启用 WAF 的内部（开发）版本应用进行扫描，以获得更完整的检测结果。",
    },
    "Insecure HTTP Usage": {
        "title_zh": "使用不安全的 HTTP 协议",
        "description_zh": "检测到 Web 应用使用 HTTP 协议，且未自动将用户重定向至 HTTPS。在部分情况下，明文 HTTP 通信可能被中间人（MitM）攻击利用，导致传输数据被窃取或篡改。",
        "repair_zh": "建议在 Web 应用中实施 HTTPS 重定向，将所有 HTTP 请求自动跳转至 HTTPS 加密连接。",
    },
    "SSL/TLS Not Implemented": {
        "title_zh": "未启用 SSL/TLS 加密",
        "description_zh": "目标站点通过未加密连接进行通信，潜在攻击者可拦截并篡改该站点发送和接收的数据，存在信息泄露风险。",
        "repair_zh": "站点应通过安全的 HTTPS 连接收发数据，部署 SSL/TLS 证书并强制使用加密通信。",
    },
    "Access-Control-Allow-Origin header with wildcard (*) value": {
        "title_zh": "Access-Control-Allow-Origin 头使用通配符（*）",
        "description_zh": "跨域资源共享（CORS）机制允许网页从其它域请求受限资源。当响应头为 Access-Control-Allow-Origin: * 时，资源允许与所有来源共享，任意网站都可以向该站点发起 XHR 请求并读取响应内容。",
        "repair_zh": "检查 Access-Control-Allow-Origin: * 是否适用于该资源/响应，建议仅对公开资源使用通配符，敏感资源应限定为受信任的域名。",
    },
    "Subresource Integrity (SRI) Not Implemented": {
        "title_zh": "未实施子资源完整性（SRI）校验",
        "description_zh": "从外部 URL 加载的脚本未实现子资源完整性（SRI）校验。SRI 使浏览器能够验证第三方资源（如来自 CDN 的脚本）未经意外篡改；若托管 CDN 被入侵或文件被篡改，攻击者可替换这些外部资源。",
        "repair_zh": "为所有从外部主机加载的脚本添加 SRI 完整性属性（integrity），通过 SRI Hash Generator 生成包含 base64 编码哈希的 script 标签。",
    },
    "[Possible] Internal IP Address Disclosure": {
        "title_zh": "（疑似）内网 IP 地址泄露",
        "description_zh": "在页面中检测到匹配内网 IPv4 地址的字符串。这些地址可能泄露内部网络的 IP 地址规划信息，攻击者可利用该信息实施进一步攻击。此发现的实际意义需要人工确认。",
        "repair_zh": "阻止该信息向用户展示，确保页面及响应中不包含内网 IP 地址等敏感信息。",
    },
}


def strip_html(text):
    """去除 HTML 标签，保留可读文本"""
    if not text:
        return ""
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</li>', '', text)
    text = re.sub(r'<li>', '- ', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = (text.replace('&quot;', '"').replace('&lt;', '<')
                .replace('&gt;', '>').replace('&amp;', '&')
                .replace('&#39;', "'"))
    return text.strip()


def parse_request(raw):
    """解析 HTTP 原始请求文本 -> 结构化字典"""
    if not raw:
        return {}
    lines = raw.split('\r\n')
    first = lines[0].split(' ', 2)
    method = first[0] if first else 'GET'
    url = first[1] if len(first) > 1 else ''
    headers = {}
    body = ""
    body_start = raw.find('\r\n\r\n')
    for line in lines[1:]:
        if not line:
            continue
        if ':' in line:
            k, v = line.split(':', 1)
            headers[k.strip()] = v.strip()
    if body_start >= 0:
        body = raw[body_start + 4:]
    return {'method': method, 'url': url, 'headers': headers, 'body': body}


def decode_response(raw):
    """Invicti 的 response 字段为 gzip 压缩后的 base64，解压还原"""
    if not raw:
        return {}
    try:
        data = base64.b64decode(raw)
        text = gzip.decompress(data).decode('utf-8', errors='replace')
    except Exception:
        text = str(raw)
    lines = text.split('\r\n')
    status_code = '200'
    if lines and ' ' in lines[0]:
        parts = lines[0].split(' ', 2)
        status_code = parts[1] if len(parts) > 1 else '200'
    headers = {}
    body = ""
    body_start = text.find('\r\n\r\n')
    for line in lines[1:]:
        if not line:
            continue
        if ':' in line:
            k, v = line.split(':', 1)
            headers[k.strip()] = v.strip()
    if body_start >= 0:
        body = text[body_start + 4:]
    return {'status_code': status_code, 'headers': headers, 'body': body}


def get_confidence(tags):
    """从 tags 中提取置信度"""
    for tag in tags or []:
        if tag.startswith('confidence.'):
            val = tag.split('.', 1)[1]
            return {'100': 'verified', '75': 'high', '50': 'medium'}.get(val, val)
    return 'unknown'


def convert_data(raw):
    """将 Acunetix/Invicti 导出数据转换为报告生成所需结构，返回 (data, tool_version)"""
    scans = raw['export']['scans']
    tool_version = DEFAULT_TOOL_VERSION
    target_urls = []
    site_summary = []
    all_vulns = []
    total = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'info': 0}
    end_dates = []

    vuln_seq = 0
    for scan in scans:
        info = scan.get('info', {})
        start_url = info.get('start_url', info.get('host', ''))
        target_urls.append(start_url)
        if info.get('end_date'):
            end_dates.append(info['end_date'])
        if info.get('build'):
            tool_version = info['build']

        vt_map = {vt['vt_id']: vt for vt in scan.get('vulnerability_types', [])}
        site_counts = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'info': 0}

        for v in scan.get('vulnerabilities', []):
            vin = v.get('info', {})
            vt = vt_map.get(vin.get('vt_id', ''), {})
            sev = SEVERITY_MAP.get(int(vt.get('severity', 0) or 0), 'info')
            site_counts[sev] += 1
            total[sev] += 1
            vuln_seq += 1

            vt_type = str(vt.get('type', ''))
            vt_name = vt.get('name', '')
            zh = VULN_ZH_MAP.get(vt_name, {})
            title = zh.get('title_zh') or vin.get('name', vt_name or '未知漏洞')
            description = zh.get('description_zh') or strip_html(vt.get('description', '')) or '无描述'
            repair = zh.get('repair_zh') or strip_html(vt.get('recommendation', '')) or '暂无修复建议'
            interactions = []
            req = parse_request(vin.get('request'))
            resp = decode_response(v.get('response'))
            if req or resp:
                interactions.append({
                    'seq': 1,
                    'label': '漏洞验证请求/响应',
                    'request': req,
                    'response': resp,
                })
            elif vin.get('details'):
                # 无请求/响应数据时（如 SSL 审计类漏洞），将扫描器检测详情写入漏洞详情
                interactions.append({
                    'seq': 1,
                    'label': '扫描器检测详情',
                    'request': {},
                    'response': {'status_code': '', 'headers': {}, 'body': strip_html(vin['details'])},
                })

            all_vulns.append({
                'vuln_id': f"VULN-{vuln_seq:03d}",
                'title': title,
                'severity': sev,
                'type': vt_type,
                'type_zh': TYPE_ZH_MAP.get(vt_type, vt_type or '其他'),
                'confidence': get_confidence(vin.get('tags')),
                'target_url': vin.get('url', start_url),
                'description': description,
                'inject_parameter': '',
                'inject_payload': '',
                'http_interactions': interactions,
                'RepairSuggestions': repair,
            })

        site_summary.append({
            'site_url': start_url,
            'critical': site_counts['critical'],
            'high': site_counts['high'],
            'medium': site_counts['medium'],
            'low': site_counts['low'],
            'total': sum(site_counts.values()),
        })

    total_count = sum(total.values())
    generated_at = max(end_dates) if end_dates else datetime.now().isoformat()

    data = {
        'report_meta': {
            'report_id': f"SR-{datetime.now().strftime('%Y%m%d')}-01",
            'generated_at': generated_at,
            'tool_version': tool_version,
            'scope': {
                'target_urls': target_urls,
                'tech_stack': [],
            },
            'site_vulnerability_summary': site_summary,
        },
        'summary': {
            'critical': total['critical'],
            'high': total['high'],
            'medium': total['medium'],
            'low': total['low'],
            'info': total['info'],
            'total': total_count,
        },
        'vulnerabilities': all_vulns,
    }
    return data, tool_version


# ========== 样式辅助函数 ==========
# 现代化配色方案
COLORS = {
    'primary': (27, 58, 92),      # Navy
    'secondary': (59, 191, 191),  # Teal
    'accent': (35, 73, 114),      # Navy-mid
    'text': (26, 45, 63),         # Dark text
    'muted': (92, 122, 146),      # Muted text
    'success': (46, 204, 113),    # Green
    'warning': (241, 196, 15),    # Yellow
    'danger': (192, 57, 43),      # Red
    'info': (52, 152, 219),       # Blue
    'bg_light': (244, 247, 250),  # Light background
    'bg_card': (255, 255, 255),   # Card background
    'border': (216, 226, 236),    # Border color
}

SEVERITY_COLORS = {
    'critical': {'bg': (253, 236, 234), 'text': (192, 57, 43), 'label': '严重'},
    'high': {'bg': (254, 243, 232), 'text': (183, 80, 10), 'label': '高危'},
    'medium': {'bg': (254, 252, 232), 'text': (154, 125, 0), 'label': '中危'},
    'low': {'bg': (234, 250, 241), 'text': (26, 138, 74), 'label': '低危'},
    'info': {'bg': (235, 245, 251), 'text': (26, 111, 160), 'label': '信息'},
}

def set_run_font(run, font_name='微软雅黑', font_size=10.5, bold=False, color=None):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def add_heading_custom(doc, text, level=1, font_name='微软雅黑', font_size=16,
                     bold=True, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    if color is None:
        color = COLORS['primary']
    # 使用内置标题样式以支持导航窗格
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = alignment

    # 清除默认样式并应用自定义样式
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size, bold, color)

    paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.5

    # 添加章节分隔线（一级标题）
    if level == 1:
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '3BBFBF')  # Teal color
        pBdr.append(bottom)
        paragraph._p.get_or_add_pPr().append(pBdr)

    return paragraph


def add_paragraph_custom(doc, text, font_name='微软雅黑', font_size=10.5,
                         bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=0.5, color=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.6
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph.paragraph_format.space_after = Pt(6)

    lines = str(text).split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            run = paragraph.add_run('\n')
            set_run_font(run, font_name, font_size)
        run = paragraph.add_run(line)
        set_run_font(run, font_name, font_size, bold, color)
    return paragraph


def add_code_block(doc, title, content, bg_color='F5F5F5'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"【{title}】")
    set_run_font(run, '仿宋', 10, True, (80, 80, 80))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6)

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), bg_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{edge}')
        margin.set(qn('w:w'), '100')
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

    p_cell = cell.paragraphs[0]
    p_cell.paragraph_format.line_spacing = 1.2
    p_cell.paragraph_format.space_after = Pt(0)

    lines = str(content).split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            run = p_cell.add_run('\n')
            set_run_font(run, 'Courier New', 9)
        run = p_cell.add_run(line)
        set_run_font(run, 'Courier New', 9, False, (50, 50, 50))
    return table


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_elm = OxmlElement(f'w:{edge}')
            edge_elm.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            edge_elm.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            edge_elm.set(qn('w:space'), '0')
            edge_elm.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcPr.append(edge_elm)


# ========== 图表生成 ==========
def generate_charts(summary, chart_path):
    # 使用现代化配色
    colors = ['#C0392B', '#E67E22', '#F1C40F', '#2ECC71', '#3498DB']
    labels = ['严重', '高危', '中危', '低危', '信息']
    sizes = [summary['critical'], summary['high'], summary['medium'],
             summary['low'], summary['info']]
    explode = (0.05, 0.03, 0.02, 0, 0)

    # 创建图表
    fig = plt.figure(figsize=(12, 5))
    gs = fig.add_gridspec(1, 2, width_ratios=[1, 1.2], wspace=0.3)

    # 左侧：环形图
    ax1 = fig.add_subplot(gs[0, 0])
    wedges, texts, autotexts = ax1.pie(sizes, explode=explode, labels=None, colors=colors,
                                        autopct='%1.1f%%', shadow=False, startangle=90,
                                        pctdistance=0.85, wedgeprops=dict(width=0.5))
    for autotext in autotexts:
        autotext.set_fontsize(9)
        autotext.set_fontweight('bold')
        autotext.set_color('white')

    # 添加图例
    ax1.legend(wedges, [f'{l} ({s})' for l, s in zip(labels, sizes)],
               title="漏洞类型",
               loc="lower right",
               bbox_to_anchor=(1.2, -0.1),
               fontsize=9)
    ax1.set_title('漏洞严重程度分布', fontsize=13, fontweight='bold', pad=15, color='#1B3A5C')

    # 右侧：柱状图
    ax2 = fig.add_subplot(gs[0, 1])
    bars = ax2.bar(labels, sizes, color=colors, edgecolor='white', linewidth=1.5, width=0.6)

    # 添加数值标签
    for bar in bars:
        height = bar.get_height()
        ax2.annotate(f'{int(height)}',
                     xy=(bar.get_x() + bar.get_width() / 2, height),
                     xytext=(0, 5), textcoords="offset points",
                     ha='center', va='bottom', fontsize=11, fontweight='bold',
                     color='#1B3A5C')

    ax2.set_title('各严重程度漏洞数量', fontsize=13, fontweight='bold', pad=15, color='#1B3A5C')
    ax2.set_ylabel('数量', fontsize=10, color='#5C7A92')
    ax2.set_xlabel('严重程度', fontsize=10, color='#5C7A92')
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['left'].set_color('#D8E2EC')
    ax2.spines['bottom'].set_color('#D8E2EC')
    ax2.tick_params(colors='#5C7A92')
    ax2.grid(axis='y', alpha=0.3, linestyle='--', color='#D8E2EC')

    plt.savefig(chart_path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    return chart_path


def add_site_summary_table(doc, site_summary, site_names=None):
    site_names = site_names or {}
    table = doc.add_table(rows=len(site_summary) + 1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['站点 URL', '站点名称', '严重', '高危', '中危', '低危', '总数']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, '仿宋', 11, True, (255, 255, 255))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E79')
        cell._tc.get_or_add_tcPr().append(shading)

    risk_styles = {
        2: {"fill": "FDECEA", "font_color": (192, 57, 43)},
        3: {"fill": "FEF3E8", "font_color": (183, 80, 10)},
        4: {"fill": "FEFCE8", "font_color": (154, 125, 0)},
        5: {"fill": "EAFAF1", "font_color": (26, 138, 74)},
        6: {"fill": "EEF2F6", "font_color": (27, 58, 92)},
    }

    for row_idx, row_data in enumerate(site_summary, 1):
        values = [
            row_data.get('site_url', ''),
            site_names.get(normalise_site_url(row_data.get('site_url', '')), ''),
            str(row_data.get('critical', 0)),
            str(row_data.get('high', 0)),
            str(row_data.get('medium', 0)),
            str(row_data.get('low', 0)),
            str(row_data.get('total', 0)),
        ]
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            font_color = risk_styles.get(col_idx, {}).get("font_color")
            set_run_font(run, '仿宋', 10.5, False, font_color)
            if col_idx in risk_styles:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), risk_styles[col_idx]["fill"])
                cell._tc.get_or_add_tcPr().append(shading)

    return table


def normalise_site_url(url):
    return (url or "").strip().rstrip("/")


def load_fingerprint_data():
    fingerprint_path = os.path.join(os.getcwd(), "workspace", "fingerprint.json")
    try:
        with open(fingerprint_path, "r", encoding="utf-8") as f:
            fingerprint_data = json.load(f)
    except Exception:
        return {}

    if isinstance(fingerprint_data, dict):
        return fingerprint_data
    return {}


def resolve_site_name_from_fingerprint(target_url, fingerprint_data, allow_single_fallback=False):
    if not isinstance(fingerprint_data, dict) or not fingerprint_data:
        return ""

    normalized_target = normalise_site_url(target_url)
    if normalized_target:
        for site_url, site_info in fingerprint_data.items():
            if normalise_site_url(site_url) != normalized_target:
                continue
            if isinstance(site_info, dict):
                return str(site_info.get("web_title", "") or "")

    if allow_single_fallback and len(fingerprint_data) == 1:
        only_site_info = next(iter(fingerprint_data.values()))
        if isinstance(only_site_info, dict):
            return str(only_site_info.get("web_title", "") or "")

    return ""


def generate_default_output_path(report_meta):
    """根据报告元数据自动生成输出文件名"""
    report_id = report_meta.get('report_id', 'UNKNOWN')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    default_name = f"report_{report_id}_{timestamp}.docx"
    # 输出到 workspace/report_result 目录
    output_dir = os.path.join(os.getcwd(), "workspace", "report_result")
    return os.path.join(output_dir, default_name)


# ========== 报告生成主逻辑 ==========
def generate_report(data, output_path, chart_path, fingerprint_data=None):
    report_meta = data['report_meta']
    summary = data['summary']
    vulns = data['vulnerabilities']
    scope = report_meta['scope']
    tool_version = report_meta.get('tool_version', DEFAULT_TOOL_VERSION)
    target_urls = scope.get('target_urls', [])
    site_summary = report_meta.get('site_vulnerability_summary', [])
    is_multi_target = len(target_urls) >= 2
    display_target = scope.get('target_url', 'N/A')
    site_name = ""
    site_names = {}
    if is_multi_target:
        display_target = f"多个目标（{len(target_urls)}个站点）"
        for row in site_summary:
            site_url = row.get('site_url', '')
            site_names[normalise_site_url(site_url)] = resolve_site_name_from_fingerprint(
                site_url,
                fingerprint_data,
                allow_single_fallback=False,
            )
    tech_stack_text = ", ".join(scope.get('tech_stack', []))
    if not is_multi_target:
        site_name = resolve_site_name_from_fingerprint(
            scope.get('target_url', ''),
            fingerprint_data,
            allow_single_fallback=True,
        )
    risk_ratio = (
        (summary['critical'] + summary['high']) / summary['total'] * 100
        if summary['total'] else 0.0
    )

    doc = Document()

    # ========== 封面页 ==========
    # 添加顶部装饰线
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '24')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '3BBFBF')  # Teal
    pBdr.append(top)
    pPr.append(pBdr)

    # 空行
    for _ in range(4):
        doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Web 应用安全")
    set_run_font(run, '微软雅黑', 32, True, COLORS['primary'])
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("漏洞扫描报告")
    set_run_font(run, '微软雅黑', 32, True, COLORS['primary'])
    p.paragraph_format.space_after = Pt(24)

    # 副标题/装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 20)
    set_run_font(run, '微软雅黑', 10, False, COLORS['secondary'])
    p.paragraph_format.space_after = Pt(36)

    # 信息表格
    info_items = [
        ("测试目标", display_target),
        ("测试时间", report_meta['generated_at'].replace('T', ' ').replace('Z', '')),
        ("测试工具", TOOL_NAME),
        ("工具版本", tool_version),
        ("漏洞总数", str(summary['total'])),
    ]
    if not is_multi_target:
        info_items.insert(1, ("站点名称", site_name))
        info_items.insert(3, ("技术栈", tech_stack_text))

    table = doc.add_table(rows=len(info_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (key, val) in enumerate(info_items):
        row = table.rows[i]
        row.height = Pt(32)

        # 左侧标签单元格
        cell_key = row.cells[0]
        cell_key.width = Inches(1.8)
        p = cell_key.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(key)
        set_run_font(run, '微软雅黑', 11, True, (255, 255, 255))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1B3A5C')  # Navy
        cell_key._tc.get_or_add_tcPr().append(shading)

        # 右侧值单元格
        cell_val = row.cells[1]
        cell_val.width = Inches(3.5)
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        set_run_font(run, '微软雅黑', 11, False, COLORS['text'])

        # 交替行背景
        if i % 2 == 0:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F4F7FA')
            cell_val._tc.get_or_add_tcPr().append(shading)

        # 设置边框
        for cell in [cell_key, cell_val]:
            set_cell_border(cell,
                top={'val':'single','sz':4,'color':'D8E2EC'},
                bottom={'val':'single','sz':4,'color':'D8E2EC'},
                left={'val':'single','sz':4,'color':'D8E2EC'},
                right={'val':'single','sz':4,'color':'D8E2EC'})

    doc.add_page_break()

    # 目录
    add_heading_custom(doc, "目  录", level=1, font_size=18,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(0, 51, 102))
    doc.add_paragraph()

    toc_items = ["一、 漏洞扫描概述", "二、 漏洞风险统计"]
    if is_multi_target:
        toc_items.append("三、 网站漏洞列表")
        detail_heading = "四、 漏洞详情"
        repair_heading = "五、 修复建议汇总"
    else:
        detail_heading = "三、 漏洞详情"
        repair_heading = "四、 修复建议汇总"
    toc_items.extend([detail_heading, repair_heading])
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(item)
        set_run_font(run, '仿宋', 14, True, (0, 51, 102))

    doc.add_page_break()

    # 第一章
    add_heading_custom(doc, "一、 漏洞扫描概述", level=1, font_size=18, color=(0, 51, 102))
    overview_lines = [
        f"本次安全测试针对目标系统 {display_target} 进行了全面的Web应用漏洞扫描。",
        f"扫描时间：{report_meta['generated_at'].replace('T', ' ').replace('Z', '')}",
        f"测试工具：{TOOL_NAME}",
        f"工具版本：{tool_version}",
    ]
    if not is_multi_target:
        if site_name:
            overview_lines.append(f"站点名称：{site_name}")
        overview_lines.append(f"目标技术栈：{tech_stack_text}")
    elif site_names:
        named_sites = [name for name in site_names.values() if name]
        if named_sites:
            overview_lines.append(f"涉及站点名称：{'；'.join(named_sites)}")
    overview_lines.extend([
        "",
        f"本次扫描共发现 {summary['total']} 个安全漏洞，其中：",
        f"- 严重（Critical）：{summary['critical']} 个",
        f"- 高危（High）：{summary['high']} 个",
        f"- 中危（Medium）：{summary['medium']} 个",
        f"- 低危（Low）：{summary['low']} 个",
        f"- 信息（Info）：{summary['info']} 个",
        "",
        f"严重及高危漏洞占比 {risk_ratio:.1f}%，建议优先修复。",
    ])
    overview_text = "\n".join(overview_lines)
    add_paragraph_custom(doc, overview_text, font_size=11)
    doc.add_page_break()

    # 第二章
    add_heading_custom(doc, "二、 漏洞风险统计", level=1, font_size=18, color=(0, 51, 102))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(5.8))
    p.paragraph_format.space_after = Pt(12)

    add_heading_custom(doc, "漏洞类型分布", level=2, font_size=14, color=(0, 64, 128))

    type_count = {}
    for v in vulns:
        t = v.get('type_zh', v.get('type', '未知'))
        type_count[t] = type_count.get(t, 0) + 1

    table = doc.add_table(rows=len(type_count)+1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['序号', '漏洞类型', '数量']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, '仿宋', 11, True, (255, 255, 255))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E79')
        cell._tc.get_or_add_tcPr().append(shading)

    for idx, (t, c) in enumerate(sorted(type_count.items(), key=lambda x: -x[1]), 1):
        row = table.rows[idx]
        for col_idx, val in enumerate([str(idx), t, str(c)]):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, '微软雅黑', 10.5)
            # 交替行背景
            if idx % 2 == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F4F7FA')
                cell._tc.get_or_add_tcPr().append(shading)

    if is_multi_target:
        doc.add_page_break()
        add_heading_custom(doc, "三、 网站漏洞列表", level=1, font_size=18)
        add_paragraph_custom(
            doc,
            "以下列表按站点维度统计严重、高危、中危、低危和总漏洞数，用于快速识别高风险站点。",
            font_size=11,
            first_line_indent=0,
        )
        add_site_summary_table(doc, site_summary, site_names=site_names)
    doc.add_page_break()

    # 漏洞详情
    add_heading_custom(doc, detail_heading, level=1, font_size=18)

    # 使用新的配色方案
    severity_names = {
        'critical': '严重', 'high': '高危', 'medium': '中危', 'low': '低危', 'info': '信息'
    }

    for idx, vuln in enumerate(vulns, 1):
        sev = vuln.get('severity', 'info').lower()
        sev_info = SEVERITY_COLORS.get(sev, {'text': (100, 100, 100), 'label': sev})
        sev_color = sev_info['text']
        sev_name = sev_info['label']

        title_text = f"{idx}. 【{sev_name}】{vuln['title']}"
        add_heading_custom(doc, title_text, level=2, font_size=14, color=sev_color)

        info_table = doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        info_table.autofit = False

        info_data = [
            ("漏洞编号", vuln['vuln_id']),
            ("风险等级", f"{sev_name} ({vuln['severity']})"),
            ("漏洞类型", vuln.get('type_zh', vuln.get('type', '未知'))),
            ("确认状态", vuln.get('confidence', 'unknown')),
            ("目标地址", vuln.get('target_url', 'N/A')),
        ]
        for i, (k, v) in enumerate(info_data):
            row = info_table.rows[i]
            cell_k = row.cells[0]
            cell_v = row.cells[1]
            cell_k.width = Inches(1.5)
            cell_v.width = Inches(4.5)

            p = cell_k.paragraphs[0]
            run = p.add_run(k)
            set_run_font(run, '仿宋', 10.5, True, (255, 255, 255))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '2E75B6')
            cell_k._tc.get_or_add_tcPr().append(shading)

            p = cell_v.paragraphs[0]
            run = p.add_run(v)
            set_run_font(run, '微软雅黑', 10.5, color=COLORS['text'])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 添加边框（使用较浅的颜色）
            for cell in [cell_k, cell_v]:
                set_cell_border(cell,
                    top={'val':'single', 'sz':4, 'color':'D8E2EC'},
                    bottom={'val':'single', 'sz':4, 'color':'D8E2EC'},
                    left={'val':'single', 'sz':4, 'color':'D8E2EC'},
                    right={'val':'single', 'sz':4, 'color':'D8E2EC'})

        add_heading_custom(doc, "漏洞描述", level=3, font_size=12)
        add_paragraph_custom(doc, vuln.get('description', '无'), font_size=10.5)

        if vuln.get('inject_parameter') or vuln.get('inject_payload'):
            add_heading_custom(doc, "漏洞参数", level=3, font_size=12)
            if vuln.get('inject_parameter'):
                add_paragraph_custom(doc, f"注入参数：{vuln['inject_parameter']}",
                                     font_size=10.5, first_line_indent=0)
            if vuln.get('inject_payload'):
                add_paragraph_custom(doc, f"测试载荷：{vuln['inject_payload']}",
                                     font_size=10.5, first_line_indent=0)

        interactions = vuln.get('http_interactions', [])
        if interactions:
            add_heading_custom(doc, "漏洞详情", level=3, font_size=12)
            for inter in interactions:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(f"► 交互 {inter['seq']}：{inter.get('label', '无标签')}")
                set_run_font(run, '微软雅黑', 10.5, True, COLORS['secondary'])

                req = inter.get('request', {})
                resp = inter.get('response', {})

                if req:
                    req_lines = [f"{req.get('method', 'GET')} {req.get('url', '')} HTTP/1.1"]
                    for h_k, h_v in req.get('headers', {}).items():
                        req_lines.append(f"{h_k}: {h_v}")
                    if req.get('body'):
                        req_lines.append("")
                        req_lines.append(req['body'])
                    add_code_block(doc, "请求", "\n".join(req_lines), 'F0F8FF')

                if resp:
                    if resp.get('status_code') or resp.get('headers'):
                        resp_lines = [f"HTTP/1.1 {resp.get('status_code', '200')} OK"]
                        for h_k, h_v in resp.get('headers', {}).items():
                            resp_lines.append(f"{h_k}: {h_v}")
                        if resp.get('body'):
                            resp_lines.append("")
                            body = resp['body']
                            if len(body) > 1500:
                                body = body[:1500] + "\n...[内容已截断]"
                            resp_lines.append(body)
                    else:
                        resp_lines = [resp.get('body', '')]
                    add_code_block(doc, "响应", "\n".join(resp_lines), 'FFF5EE')

        add_heading_custom(doc, "修复建议", level=3, font_size=12)
        repair = vuln.get('RepairSuggestions', '暂无修复建议')
        add_paragraph_custom(doc, repair, font_size=10.5)

        if idx < len(vulns):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 50)
            set_run_font(run, '微软雅黑', 8, False, COLORS['border'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 修复建议汇总
    add_heading_custom(doc, repair_heading, level=1, font_size=18)
    repair_summary = """根据本次扫描结果，建议从以下几个方面进行系统性修复：

1. SQL注入漏洞修复（最高优先级）
   • 对所有数据库交互点实施参数化查询（Prepared Statements），彻底杜绝字符串拼接SQL语句。
   • 在应用层实施输入验证和白名单过滤，拒绝异常字符。
   • 对数据库账户实施最小权限原则，避免使用root/dba账户连接数据库。
   • 部署WAF（Web应用防火墙）拦截常见SQL注入payload。
   • 开启数据库慢查询日志监控异常请求。

2. 认证与会话安全
   • 对所有敏感接口实施严格的认证检查，确保未授权请求被拦截。
   • 使用统一的认证中间件/过滤器确保所有路由受保护。
   • 实施基于角色的访问控制(RBAC)。
   • Cookie必须设置HttpOnly、Secure、SameSite标志。
   • 登录后重新生成会话标识符，防止会话固定攻击。

3. 信息泄露治理
   • 配置Web服务器隐藏版本信息（移除Server头、X-Powered-By头）。
   • 确保错误页面不泄露堆栈跟踪、代码路径和SQL语句。
   • 移除生产环境中的调试端点（如setup-db.php、.git目录等）。
   • 审查所有API响应，移除不应暴露的敏感字段（密码哈希、内部ID、密钥等）。

4. CSRF防护
   • 确保所有状态修改请求都携带有效的CSRF Token。
   • 验证CSRF Token在服务端的唯一性和有效性。
   • 检查Cookie的SameSite标志位设置为Strict或Lax。
   • 对关键操作添加二次验证（如密码确认、短信验证码）。

5. 安全响应头加固
   • 添加X-Frame-Options: DENY或SAMEORIGIN，防止Clickjacking。
   • 添加Content-Security-Policy头，限制资源加载来源。
   • 添加X-Content-Type-Options: nosniff，防止MIME嗅探。
   • 部署HTTPS并添加Strict-Transport-Security头。

6. 代码与配置安全
   • 使用.gitignore并配置Web服务器禁止访问.git目录。
   • 将配置文件（如.inc文件）移出Web根目录或禁止直接访问。
   • 引入DTO模式，将内部数据模型与API响应解耦，仅返回前端必需的最小字段集。"""
    add_paragraph_custom(doc, repair_summary, font_size=11, first_line_indent=0)

    doc.save(output_path)
    return output_path


def print_usage():
    print("用法: python awvs_json_word.py <report.json> [output.docx]")
    print("  report.json  - Acunetix/Invicti 导出的漏洞扫描 JSON 文件路径（必须）")
    print("  output.docx  - 输出的 Word 文件路径（可选）")
    print()
    print("示例:")
    print("  python awvs_json_word.py report.json")
    print("  python awvs_json_word.py report.json my_report.docx")
    sys.exit(1)


def main():
    if len(sys.argv) < 2 or sys.argv[1] in ('-h', '--help', '/?'):
        print_usage()

    input_path = os.path.abspath(sys.argv[1])
    if not os.path.exists(input_path):
        print(f"[错误] 输入文件不存在: {input_path}")
        sys.exit(1)

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            raw = json.load(f)
    except json.JSONDecodeError as e:
        print(f"[错误] JSON 解析失败: {e}")
        sys.exit(1)

    if 'export' not in raw or 'scans' not in raw['export']:
        print("[错误] JSON 缺少 export.scans 字段，不是 Acunetix/Invicti 导出格式")
        sys.exit(1)

    data, tool_version = convert_data(raw)

    # 确保输出目录存在
    output_dir = os.path.join(os.getcwd(), "workspace", "report_result")
    os.makedirs(output_dir, exist_ok=True)

    if len(sys.argv) >= 3:
        output_path = os.path.join(output_dir, os.path.basename(sys.argv[2]))
        if not output_path.lower().endswith('.docx'):
            output_path += '.docx'
    else:
        output_path = generate_default_output_path(data['report_meta'])

    chart_path = generate_charts(data['summary'], os.path.join(output_dir, 'vuln_chart.png'))

    print(f"[信息] 输入文件: {input_path}")
    print(f"[信息] 测试工具: {TOOL_NAME} (版本 {tool_version})")
    print(f"[信息] 正在生成报告，共 {len(data['vulnerabilities'])} 个漏洞...")

    final_path = generate_report(data, output_path, chart_path, fingerprint_data=load_fingerprint_data())

    print(f"[成功] 报告已生成: {final_path}")
    print(f"[成功] 统计图表: {chart_path}")


if __name__ == '__main__':
    main()
